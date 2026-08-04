"""
meetings/MMblueprint.py
Meeting Minutes app - Flask blueprint.

Mount in app.py with:
    from meetings.MMblueprint import mm_bp
    app.register_blueprint(mm_bp, url_prefix="/mm")

Environment variables used:
    SUPABASE_URL, SUPABASE_KEY   - same project as SMC
    MM_PASSWORD                  - login password (default: sandle2026)
    ASSEMBLYAI_KEY               - transcription (optional until phase 1b)
    ANTHROPIC_API_KEY            - AI minute drafting
"""
from flask import Blueprint, render_template, request, jsonify, redirect, url_for, make_response, Response
from functools import wraps
import os, hashlib, json, base64, datetime

mm_bp = Blueprint("mm", __name__, template_folder="templates")

from supabase import create_client
_url = os.environ.get("SUPABASE_URL", "")
_key = os.environ.get("SUPABASE_KEY", "")
supabase = create_client(_url, _key) if _url and _key else None

MM_PASSWORD = os.environ.get("MM_PASSWORD", "sandle2026")
COOKIE_NAME = "mm_auth"
OWNER_ID    = "aaron"          # single user for now; becomes per-session later
BUCKET      = "mm-files"       # Supabase Storage bucket for audio + documents


# ══════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════
def _hash():
    return hashlib.sha256(MM_PASSWORD.encode()).hexdigest()


def ok(payload=None):
    return jsonify(payload if payload is not None else {"ok": True})


def err(e):
    return jsonify({"ok": False, "error": str(e)}), 500


def _require_auth(fn):
    @wraps(fn)
    def wrapper(*a, **kw):
        if request.cookies.get(COOKIE_NAME) != _hash():
            if request.path.startswith("/mm/api/"):
                return jsonify({"ok": False, "error": "unauthorised"}), 401
            return redirect(url_for("mm.login"))
        return fn(*a, **kw)
    return wrapper


# ══════════════════════════════════════════════════════════════════
# Auth + app shell
# ══════════════════════════════════════════════════════════════════
def _read_template(name):
    """Read an HTML file from templates/mm/ directly, bypassing Jinja2.
    Avoids any template-resolution issues and keeps React braces untouched."""
    here = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(here, "templates", "mm", name)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


@mm_bp.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        if request.form.get("password", "") == MM_PASSWORD:
            resp = make_response(redirect(url_for("mm.app_page")))
            resp.set_cookie(COOKIE_NAME, _hash(), max_age=60 * 60 * 24 * 90,
                            httponly=True, samesite="Lax", path="/")
            return resp
        error = "Incorrect password"
    try:
        html = _read_template("login.html")
    except FileNotFoundError:
        return Response(
            "<h2>login.html not found</h2>"
            "<p>Expected at meetings/templates/mm/login.html</p>"
            "<p>Visit <a href='/mm/health'>/mm/health</a> for details.</p>",
            mimetype="text/html", status=500)
    # Simple placeholder substitution instead of Jinja2
    block = ('<div class="err">' + error + "</div>") if error else ""
    html = html.replace("{% if error %}<div class=\"err\">{{ error }}</div>{% endif %}", block)
    return Response(html, mimetype="text/html")


@mm_bp.route("/logout")
def logout():
    resp = make_response(redirect(url_for("mm.login")))
    resp.delete_cookie(COOKIE_NAME, path="/")
    return resp


@mm_bp.route("/")
@mm_bp.route("/app")
@_require_auth
def app_page():
    """Serve the SPA directly as text/html, bypassing Jinja2 so React
    template braces are never touched."""
    try:
        return Response(_read_template("MMapp.html"), mimetype="text/html")
    except FileNotFoundError:
        return Response(
            "<h2>MMapp.html not found</h2>"
            "<p>Expected at meetings/templates/mm/MMapp.html</p>"
            "<p>Visit <a href='/mm/health'>/mm/health</a> for details.</p>",
            mimetype="text/html", status=500)


@mm_bp.route("/health")
def health():
    """Diagnostic - no auth required. Reports exactly what is on disk and
    whether the database connection works."""
    here = os.path.dirname(os.path.abspath(__file__))
    tpl_dir = os.path.join(here, "templates", "mm")

    def check(p):
        return {"path": p, "exists": os.path.exists(p),
                "size": os.path.getsize(p) if os.path.exists(p) else 0}

    info = {
        "blueprint_file": __file__,
        "blueprint_dir":  here,
        "dir_contents":   sorted(os.listdir(here)) if os.path.exists(here) else "MISSING",
        "templates_mm_exists": os.path.exists(tpl_dir),
        "templates_mm_contents": sorted(os.listdir(tpl_dir)) if os.path.exists(tpl_dir) else "MISSING",
        "MMapp_html":  check(os.path.join(tpl_dir, "MMapp.html")),
        "login_html":  check(os.path.join(tpl_dir, "login.html")),
        "supabase_connected": supabase is not None,
        "env": {
            "SUPABASE_URL_set":   bool(os.environ.get("SUPABASE_URL")),
            "SUPABASE_KEY_set":   bool(os.environ.get("SUPABASE_KEY")),
            "MM_PASSWORD_set":    bool(os.environ.get("MM_PASSWORD")),
            "ASSEMBLYAI_KEY_set": bool(os.environ.get("ASSEMBLYAI_KEY")),
            "ANTHROPIC_API_KEY_set": bool(os.environ.get("ANTHROPIC_API_KEY")),
        },
    }
    # Test a live query
    try:
        r = supabase.table("mm_templates").select("id,name").limit(5).execute()
        info["db_test"] = {"ok": True, "templates_found": len(r.data or []),
                           "names": [t["name"] for t in (r.data or [])]}
    except Exception as ex:
        info["db_test"] = {"ok": False, "error": str(ex)}

    return jsonify(info)


# ══════════════════════════════════════════════════════════════════
# Projects
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/projects", methods=["GET"])
@_require_auth
def api_projects():
    try:
        r = supabase.table("mm_projects").select("*") \
            .eq("owner_id", OWNER_ID).order("name").execute()
        return jsonify(r.data or [])
    except Exception as e:
        return err(e)


@mm_bp.route("/api/projects", methods=["POST"])
@_require_auth
def api_projects_create():
    d = request.get_json() or {}
    try:
        row = {
            "owner_id":   OWNER_ID,
            "name":       d.get("name", "Untitled Site"),
            "project_no": d.get("project_no"),
            "address":    d.get("address"),
            "smc_project_id": d.get("smc_project_id"),
        }
        r = supabase.table("mm_projects").insert(row).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/projects/<pid>", methods=["PATCH"])
@_require_auth
def api_projects_patch(pid):
    d = request.get_json() or {}
    allowed = {"name", "project_no", "address", "smc_project_id", "active"}
    upd = {k: v for k, v in d.items() if k in allowed}
    try:
        supabase.table("mm_projects").update(upd).eq("id", pid).execute()
        return ok()
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# People
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/people", methods=["GET"])
@_require_auth
def api_people():
    try:
        r = supabase.table("mm_people").select("*") \
            .eq("owner_id", OWNER_ID).order("name").execute()
        return jsonify(r.data or [])
    except Exception as e:
        return err(e)


@mm_bp.route("/api/people", methods=["POST"])
@_require_auth
def api_people_create():
    d = request.get_json() or {}
    name = d.get("name", "").strip()
    if not name:
        return jsonify({"ok": False, "error": "name required"}), 400
    # Auto-derive initials if not supplied
    initials = d.get("initials") or "".join(p[0].upper() for p in name.split()[:2])
    try:
        row = {
            "owner_id": OWNER_ID, "name": name, "initials": initials,
            "email": d.get("email"), "company": d.get("company"),
            "job_role": d.get("job_role"),
        }
        r = supabase.table("mm_people").insert(row).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/people/<pid>", methods=["PATCH", "DELETE"])
@_require_auth
def api_people_modify(pid):
    try:
        if request.method == "DELETE":
            supabase.table("mm_people").update({"active": False}).eq("id", pid).execute()
            return ok()
        d = request.get_json() or {}
        allowed = {"name", "initials", "email", "company", "job_role", "active"}
        upd = {k: v for k, v in d.items() if k in allowed}
        supabase.table("mm_people").update(upd).eq("id", pid).execute()
        return ok()
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Templates + their agenda items
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/templates", methods=["GET"])
@_require_auth
def api_templates():
    try:
        t = supabase.table("mm_templates").select("*") \
            .eq("owner_id", OWNER_ID).order("name").execute()
        i = supabase.table("mm_template_items").select("*") \
            .order("sort_order").execute()
        items_by_tpl = {}
        for row in (i.data or []):
            items_by_tpl.setdefault(row["template_id"], []).append(row)
        out = []
        for tpl in (t.data or []):
            tpl["items"] = items_by_tpl.get(tpl["id"], [])
            out.append(tpl)
        return jsonify(out)
    except Exception as e:
        return err(e)


@mm_bp.route("/api/templates", methods=["POST"])
@_require_auth
def api_templates_create():
    d = request.get_json() or {}
    try:
        row = {
            "owner_id": OWNER_ID,
            "name": d.get("name", "New Template"),
            "doc_title": d.get("doc_title"),
            "has_action_column": d.get("has_action_column", True),
            "has_programme_review": d.get("has_programme_review", False),
            "project_id": d.get("project_id"),
        }
        r = supabase.table("mm_templates").insert(row).execute()
        tpl = r.data[0]
        # Copy items if cloning from another template
        src = d.get("clone_from")
        if src:
            si = supabase.table("mm_template_items").select("*") \
                .eq("template_id", src).order("sort_order").execute()
            for it in (si.data or []):
                supabase.table("mm_template_items").insert({
                    "template_id": tpl["id"], "item_no": it["item_no"],
                    "title": it["title"], "parent_no": it.get("parent_no"),
                    "sort_order": it["sort_order"],
                }).execute()
        return jsonify(tpl)
    except Exception as e:
        return err(e)


@mm_bp.route("/api/templates/<tid>/items", methods=["POST"])
@_require_auth
def api_template_items_replace(tid):
    """Replace the whole item list for a template in one call."""
    items = (request.get_json() or {}).get("items", [])
    try:
        supabase.table("mm_template_items").delete().eq("template_id", tid).execute()
        for idx, it in enumerate(items):
            supabase.table("mm_template_items").insert({
                "template_id": tid,
                "item_no":  str(it.get("item_no", idx + 1)),
                "title":    it.get("title", ""),
                "parent_no": it.get("parent_no"),
                "sort_order": it.get("sort_order", idx),
            }).execute()
        return ok()
    except Exception as e:
        return err(e)


@mm_bp.route("/api/templates/<tid>", methods=["PATCH", "DELETE"])
@_require_auth
def api_templates_modify(tid):
    try:
        if request.method == "DELETE":
            supabase.table("mm_templates").delete().eq("id", tid).execute()
            return ok()
        d = request.get_json() or {}
        allowed = {"name", "doc_title", "has_action_column",
                   "has_programme_review", "project_id"}
        upd = {k: v for k, v in d.items() if k in allowed}
        supabase.table("mm_templates").update(upd).eq("id", tid).execute()
        return ok()
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Meeting series
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/series", methods=["GET", "POST"])
@_require_auth
def api_series():
    try:
        if request.method == "GET":
            r = supabase.table("mm_series").select("*") \
                .eq("owner_id", OWNER_ID).order("name").execute()
            return jsonify(r.data or [])
        d = request.get_json() or {}
        row = {
            "owner_id":   OWNER_ID,
            "project_id": d.get("project_id"),
            "template_id": d.get("template_id"),
            "name":       d.get("name", "New Series"),
            "frequency":  d.get("frequency", "weekly"),
        }
        r = supabase.table("mm_series").insert(row).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Meetings
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/meetings", methods=["GET"])
@_require_auth
def api_meetings():
    try:
        q = supabase.table("mm_meetings").select("*").eq("owner_id", OWNER_ID)
        pid = request.args.get("project_id")
        if pid:
            q = q.eq("project_id", pid)
        r = q.order("meeting_date", desc=True).limit(300).execute()
        return jsonify(r.data or [])
    except Exception as e:
        return err(e)


@mm_bp.route("/api/meetings/<mid>", methods=["GET"])
@_require_auth
def api_meeting_detail(mid):
    """Full meeting: header, agenda items, actions, and the previous
    meeting's open actions for Matters Arising."""
    try:
        m = supabase.table("mm_meetings").select("*").eq("id", mid).single().execute()
        meeting = m.data
        items = supabase.table("mm_meeting_items").select("*") \
            .eq("meeting_id", mid).order("sort_order").execute()
        actions = supabase.table("mm_actions").select("*") \
            .eq("meeting_id", mid).order("created_at").execute()

        # Open actions carried in from earlier meetings in the same series
        carried = []
        if meeting.get("series_id"):
            prev = supabase.table("mm_meetings").select("id") \
                .eq("series_id", meeting["series_id"]) \
                .lt("meeting_date", meeting["meeting_date"]) \
                .order("meeting_date", desc=True).limit(1).execute()
            if prev.data:
                c = supabase.table("mm_actions").select("*") \
                    .eq("meeting_id", prev.data[0]["id"]) \
                    .neq("status", "complete").execute()
                carried = c.data or []

        return jsonify({
            "meeting": meeting,
            "items":   items.data or [],
            "actions": actions.data or [],
            "carried_actions": carried,
        })
    except Exception as e:
        return err(e)


@mm_bp.route("/api/meetings", methods=["POST"])
@_require_auth
def api_meetings_create():
    """Create a meeting and SNAPSHOT the template's agenda items onto it,
    so editing the template later never rewrites historic minutes."""
    d = request.get_json() or {}
    try:
        row = {
            "owner_id":    OWNER_ID,
            "series_id":   d.get("series_id"),
            "project_id":  d.get("project_id"),
            "template_id": d.get("template_id"),
            "meeting_date": d.get("meeting_date") or datetime.date.today().isoformat(),
            "meeting_time": d.get("meeting_time"),
            "title":       d.get("title"),
            "location":    d.get("location"),
            "attendees":   d.get("attendees"),
            "apologies":   d.get("apologies"),
            "status":      "draft",
        }
        r = supabase.table("mm_meetings").insert(row).execute()
        meeting = r.data[0]

        # Snapshot agenda items from the template
        if d.get("template_id"):
            ti = supabase.table("mm_template_items").select("*") \
                .eq("template_id", d["template_id"]).order("sort_order").execute()
            for it in (ti.data or []):
                supabase.table("mm_meeting_items").insert({
                    "meeting_id": meeting["id"],
                    "item_no":    it["item_no"],
                    "title":      it["title"],
                    "parent_no":  it.get("parent_no"),
                    "sort_order": it["sort_order"],
                    "content":    "",
                }).execute()

        # Auto-carry open actions from the previous meeting in this series
        if d.get("series_id"):
            prev = supabase.table("mm_meetings").select("id") \
                .eq("series_id", d["series_id"]) \
                .lt("meeting_date", row["meeting_date"]) \
                .order("meeting_date", desc=True).limit(1).execute()
            if prev.data:
                open_acts = supabase.table("mm_actions").select("*") \
                    .eq("meeting_id", prev.data[0]["id"]) \
                    .neq("status", "complete").execute()
                for a in (open_acts.data or []):
                    supabase.table("mm_actions").insert({
                        "owner_id":      OWNER_ID,
                        "meeting_id":    meeting["id"],
                        "project_id":    a.get("project_id"),
                        "description":   a["description"],
                        "assigned_to":   a.get("assigned_to"),
                        "assigned_text": a.get("assigned_text"),
                        "due_date":      a.get("due_date"),
                        "status":        a.get("status", "open"),
                        "priority":      a.get("priority", "normal"),
                        "carried_from":  a["id"],
                    }).execute()

        return jsonify(meeting)
    except Exception as e:
        return err(e)


@mm_bp.route("/api/meetings/<mid>", methods=["PATCH", "DELETE"])
@_require_auth
def api_meetings_modify(mid):
    try:
        if request.method == "DELETE":
            supabase.table("mm_meetings").delete().eq("id", mid).execute()
            return ok()
        d = request.get_json() or {}
        allowed = {
            "title", "meeting_date", "meeting_time", "location", "status",
            "attendees", "apologies", "transcript", "audio_url",
            "transcribe_status", "series_id", "template_id", "project_id",
            "prog_baseline_duration", "prog_prior_duration",
            "prog_current_duration", "prog_baseline_completion",
            "prog_anticipated_completion",
        }
        upd = {k: v for k, v in d.items() if k in allowed}
        upd["updated_at"] = datetime.datetime.utcnow().isoformat()
        supabase.table("mm_meetings").update(upd).eq("id", mid).execute()
        return ok()
    except Exception as e:
        return err(e)


@mm_bp.route("/api/meeting-items/<iid>", methods=["PATCH"])
@_require_auth
def api_meeting_item_patch(iid):
    d = request.get_json() or {}
    allowed = {"content", "title", "item_no", "sort_order"}
    upd = {k: v for k, v in d.items() if k in allowed}
    try:
        supabase.table("mm_meeting_items").update(upd).eq("id", iid).execute()
        return ok()
    except Exception as e:
        return err(e)


@mm_bp.route("/api/meetings/<mid>/items", methods=["POST"])
@_require_auth
def api_meeting_item_add(mid):
    d = request.get_json() or {}
    try:
        r = supabase.table("mm_meeting_items").insert({
            "meeting_id": mid,
            "item_no":   str(d.get("item_no", "")),
            "title":     d.get("title", "New Item"),
            "parent_no": d.get("parent_no"),
            "sort_order": d.get("sort_order", 999),
            "content":   d.get("content", ""),
        }).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Actions - the task database
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/actions", methods=["GET"])
@_require_auth
def api_actions():
    """All actions with joined project / meeting / person context, so the
    overview page can sort by any dimension without extra lookups."""
    try:
        q = supabase.table("mm_actions").select("*").eq("owner_id", OWNER_ID)
        if request.args.get("status"):
            q = q.eq("status", request.args["status"])
        if request.args.get("project_id"):
            q = q.eq("project_id", request.args["project_id"])
        if request.args.get("assigned_to"):
            q = q.eq("assigned_to", request.args["assigned_to"])
        r = q.order("due_date").limit(1000).execute()
        actions = r.data or []

        # Enrich with names so the UI needs one request, not four
        projects = {p["id"]: p for p in (supabase.table("mm_projects")
                    .select("id,name,project_no").execute().data or [])}
        people = {p["id"]: p for p in (supabase.table("mm_people")
                  .select("id,name,initials,email").execute().data or [])}
        meetings = {m["id"]: m for m in (supabase.table("mm_meetings")
                    .select("id,title,meeting_date").execute().data or [])}

        for a in actions:
            pr = projects.get(a.get("project_id")) or {}
            pe = people.get(a.get("assigned_to")) or {}
            me = meetings.get(a.get("meeting_id")) or {}
            a["project_name"]  = pr.get("name")
            a["project_no"]    = pr.get("project_no")
            a["person_name"]   = pe.get("name") or a.get("assigned_text")
            a["person_email"]  = pe.get("email")
            a["person_initials"] = pe.get("initials")
            a["meeting_title"] = me.get("title")
            a["meeting_date"]  = me.get("meeting_date")
        return jsonify(actions)
    except Exception as e:
        return err(e)


@mm_bp.route("/api/actions", methods=["POST"])
@_require_auth
def api_actions_create():
    d = request.get_json() or {}
    try:
        row = {
            "owner_id":       OWNER_ID,
            "meeting_id":     d.get("meeting_id"),
            "meeting_item_id": d.get("meeting_item_id"),
            "project_id":     d.get("project_id"),
            "description":    d.get("description", ""),
            "assigned_to":    d.get("assigned_to"),
            "assigned_text":  d.get("assigned_text"),
            "due_date":       d.get("due_date"),
            "status":         d.get("status", "open"),
            "priority":       d.get("priority", "normal"),
        }
        r = supabase.table("mm_actions").insert(row).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/actions/<aid>", methods=["PATCH", "DELETE"])
@_require_auth
def api_actions_modify(aid):
    try:
        if request.method == "DELETE":
            supabase.table("mm_actions").delete().eq("id", aid).execute()
            return ok()
        d = request.get_json() or {}
        allowed = {"description", "assigned_to", "assigned_text", "due_date",
                   "status", "priority", "completed_date", "project_id"}
        upd = {k: v for k, v in d.items() if k in allowed}
        if upd.get("status") == "complete" and "completed_date" not in upd:
            upd["completed_date"] = datetime.date.today().isoformat()
        supabase.table("mm_actions").update(upd).eq("id", aid).execute()
        return ok()
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Programme Review - auto-fill from SMC
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/programme-review/<project_id>", methods=["GET"])
@_require_auth
def api_programme_review(project_id):
    """Derive the Programme Review box from SMC data where the site is linked.
    Baseline completion = latest CRC deadline across all plots.
    Anticipated completion = latest scheduled CRC from the live programme."""
    try:
        p = supabase.table("mm_projects").select("smc_project_id") \
            .eq("id", project_id).single().execute()
        smc_id = (p.data or {}).get("smc_project_id")
        if not smc_id:
            return jsonify({"linked": False})

        plots = supabase.table("smc_plots") \
            .select("plot_number,crc_deadline_date") \
            .eq("project_id", smc_id).execute()
        dates = [r["crc_deadline_date"] for r in (plots.data or [])
                 if r.get("crc_deadline_date")]
        if not dates:
            return jsonify({"linked": True, "has_data": False})

        latest = max(dates)
        earliest = min(dates)
        d_late = datetime.date.fromisoformat(latest)
        d_early = datetime.date.fromisoformat(earliest)
        weeks = round((d_late - d_early).days / 7)

        return jsonify({
            "linked": True,
            "has_data": True,
            "plot_count": len(plots.data or []),
            "baseline_completion": d_late.strftime("%d %b %Y"),
            "anticipated_completion": d_late.strftime("%d %b %Y"),
            "baseline_duration": f"{weeks} weeks",
            "first_completion": d_early.strftime("%d %b %Y"),
        })
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# AI - draft minutes from transcript, extract actions
# ══════════════════════════════════════════════════════════════════
def _claude(prompt, max_tokens=4000):
    import anthropic
    client = anthropic.Anthropic()
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    return msg.content[0].text if msg.content else ""


@mm_bp.route("/api/ai/draft-minutes/<mid>", methods=["POST"])
@_require_auth
def api_ai_draft_minutes(mid):
    """Classify transcript content into the meeting's fixed agenda headings.

    This is deliberately a CLASSIFICATION task against a known taxonomy
    rather than open-ended summarisation - far more reliable, and the
    output always matches the house template.
    """
    try:
        m = supabase.table("mm_meetings").select("*").eq("id", mid).single().execute()
        meeting = m.data
        transcript = (request.get_json() or {}).get("transcript") or meeting.get("transcript")
        if not transcript:
            return jsonify({"ok": False, "error": "No transcript available"}), 400

        items = supabase.table("mm_meeting_items").select("*") \
            .eq("meeting_id", mid).order("sort_order").execute()
        agenda = items.data or []
        agenda_list = "\n".join(f'{i["item_no"]}. {i["title"]}' for i in agenda)

        prompt = f"""You are minuting a UK construction site meeting for Pennyfarthing Homes.

Below is the meeting transcript, followed by the FIXED agenda headings this
meeting must be minuted against.

Your task: assign each piece of discussion to the correct agenda heading and
write concise, professional minute text for each. Use the language and tone of
UK construction site minutes - factual, brief, no filler. Where nothing was
discussed under a heading, return an empty string for it.

Also extract every ACTION agreed: who is responsible and by when.

TRANSCRIPT:
{transcript[:60000]}

AGENDA HEADINGS:
{agenda_list}

Return ONLY valid JSON, no markdown fences, in exactly this shape:
{{
  "items": [
    {{"item_no": "1", "content": "minute text here"}}
  ],
  "actions": [
    {{"description": "what must be done", "assigned_text": "person or company",
      "due_date": "YYYY-MM-DD or null", "item_no": "which agenda item it came from",
      "priority": "low|normal|high|critical"}}
  ],
  "attendees_detected": ["names heard in the transcript"]
}}"""

        raw = _claude(prompt, max_tokens=8000)
        clean = raw.replace("```json", "").replace("```", "").strip()
        data = json.loads(clean)

        # Write minute text into the snapshotted agenda items
        by_no = {i["item_no"]: i for i in agenda}
        for it in data.get("items", []):
            target = by_no.get(str(it.get("item_no")))
            if target and it.get("content"):
                supabase.table("mm_meeting_items") \
                    .update({"content": it["content"]}) \
                    .eq("id", target["id"]).execute()

        # Create actions
        created = []
        for a in data.get("actions", []):
            item = by_no.get(str(a.get("item_no")))
            row = {
                "owner_id":       OWNER_ID,
                "meeting_id":     mid,
                "meeting_item_id": item["id"] if item else None,
                "project_id":     meeting.get("project_id"),
                "description":    a.get("description", ""),
                "assigned_text":  a.get("assigned_text"),
                "due_date":       a.get("due_date") or None,
                "priority":       a.get("priority", "normal"),
                "status":         "open",
            }
            r = supabase.table("mm_actions").insert(row).execute()
            if r.data:
                created.append(r.data[0])

        return jsonify({
            "ok": True,
            "items_filled": len(data.get("items", [])),
            "actions_created": len(created),
            "attendees_detected": data.get("attendees_detected", []),
        })
    except Exception as e:
        return err(e)


@mm_bp.route("/api/ai/process-inbox", methods=["POST"])
@_require_auth
def api_ai_process_inbox():
    """Take dumped text (email, message, note) and pull out the key information:
    updates, actions, dates, issues."""
    d = request.get_json() or {}
    content = d.get("content", "")
    project_id = d.get("project_id")
    if not content.strip():
        return jsonify({"ok": False, "error": "content required"}), 400
    try:
        prompt = f"""Extract the key information from this construction site
communication (email, message or note).

CONTENT:
{content[:20000]}

Return ONLY valid JSON, no markdown fences:
{{
  "summary": "2-3 sentence summary of what this is about",
  "actions": [
    {{"description": "...", "assigned_text": "who", "due_date": "YYYY-MM-DD or null",
      "priority": "low|normal|high|critical"}}
  ],
  "key_dates": [{{"date": "YYYY-MM-DD", "what": "..."}}],
  "issues_raised": ["any problems, risks or safety matters mentioned"],
  "suggested_agenda_item": "which meeting agenda heading this belongs under"
}}"""
        raw = _claude(prompt, max_tokens=3000)
        clean = raw.replace("```json", "").replace("```", "").strip()
        data = json.loads(clean)

        inbox = supabase.table("mm_inbox").insert({
            "owner_id":    OWNER_ID,
            "project_id":  project_id,
            "source_type": d.get("source_type", "note"),
            "raw_content": content,
            "ai_summary":  data.get("summary"),
            "processed":   True,
        }).execute()

        # Optionally create the actions straight away
        if d.get("create_actions"):
            for a in data.get("actions", []):
                supabase.table("mm_actions").insert({
                    "owner_id":      OWNER_ID,
                    "project_id":    project_id,
                    "description":   a.get("description", ""),
                    "assigned_text": a.get("assigned_text"),
                    "due_date":      a.get("due_date") or None,
                    "priority":      a.get("priority", "normal"),
                    "status":        "open",
                }).execute()

        return jsonify({"ok": True, "extracted": data,
                        "inbox_id": inbox.data[0]["id"] if inbox.data else None})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/ai/ask", methods=["POST"])
@_require_auth
def api_ai_ask():
    """Ask a question across the meeting archive. Phase 1 uses keyword
    retrieval; swapped for pgvector similarity in phase 3."""
    d = request.get_json() or {}
    question = d.get("question", "")
    project_id = d.get("project_id")
    if not question.strip():
        return jsonify({"ok": False, "error": "question required"}), 400
    try:
        q = supabase.table("mm_meeting_items").select("*")
        items = q.limit(500).execute().data or []
        meetings = {m["id"]: m for m in (supabase.table("mm_meetings")
                    .select("id,title,meeting_date,project_id").execute().data or [])}

        context_parts = []
        for it in items:
            if not it.get("content"):
                continue
            mt = meetings.get(it["meeting_id"]) or {}
            if project_id and mt.get("project_id") != project_id:
                continue
            context_parts.append(
                f'[{mt.get("meeting_date","?")} | {mt.get("title","Meeting")} | '
                f'{it["item_no"]}. {it["title"]}]\n{it["content"]}'
            )
        context = "\n\n".join(context_parts[:150])

        prompt = f"""You are answering a question about construction site meeting records.

MEETING RECORDS:
{context[:80000]}

QUESTION: {question}

Answer directly and factually from the records above. Cite the meeting date and
agenda item where relevant. If the records do not contain the answer, say so
plainly rather than guessing."""
        answer = _claude(prompt, max_tokens=2000)
        return jsonify({"ok": True, "answer": answer, "sources_searched": len(context_parts)})
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Transcription (AssemblyAI)
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/transcribe/<mid>", methods=["POST"])
@_require_auth
def api_transcribe(mid):
    """Upload audio to AssemblyAI and start transcription with diarization.
    Returns immediately with a job id; poll /api/transcribe-status."""
    key = os.environ.get("ASSEMBLYAI_KEY", "")
    if not key:
        return jsonify({"ok": False, "error": "ASSEMBLYAI_KEY not configured"}), 400
    d = request.get_json() or {}
    audio_b64 = d.get("audio_b64", "")
    if not audio_b64:
        return jsonify({"ok": False, "error": "audio_b64 required"}), 400
    try:
        import requests
        audio_bytes = base64.b64decode(audio_b64)

        up = requests.post(
            "https://api.assemblyai.com/v2/upload",
            headers={"authorization": key},
            data=audio_bytes, timeout=300,
        )
        audio_url = up.json()["upload_url"]

        job = requests.post(
            "https://api.assemblyai.com/v2/transcript",
            headers={"authorization": key, "content-type": "application/json"},
            json={
                "audio_url": audio_url,
                "speaker_labels": True,      # diarization - essential for actions
                "language_code": "en_uk",
                "punctuate": True,
                "format_text": True,
            }, timeout=60,
        )
        jid = job.json()["id"]
        supabase.table("mm_meetings").update({
            "transcribe_status": "processing", "audio_url": jid,
        }).eq("id", mid).execute()
        return jsonify({"ok": True, "job_id": jid})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/transcribe-status/<mid>", methods=["GET"])
@_require_auth
def api_transcribe_status(mid):
    key = os.environ.get("ASSEMBLYAI_KEY", "")
    try:
        m = supabase.table("mm_meetings").select("audio_url,transcribe_status") \
            .eq("id", mid).single().execute()
        jid = (m.data or {}).get("audio_url")
        if not jid:
            return jsonify({"status": "none"})
        import requests
        r = requests.get(f"https://api.assemblyai.com/v2/transcript/{jid}",
                         headers={"authorization": key}, timeout=30).json()
        status = r.get("status")
        if status == "completed":
            # Build a readable diarized transcript
            lines = []
            for u in (r.get("utterances") or []):
                lines.append(f'Speaker {u["speaker"]}: {u["text"]}')
            text = "\n".join(lines) if lines else r.get("text", "")
            supabase.table("mm_meetings").update({
                "transcript": text,
                "transcript_json": r.get("utterances"),
                "transcribe_status": "done",
            }).eq("id", mid).execute()
            return jsonify({"status": "done", "transcript": text})
        if status == "error":
            supabase.table("mm_meetings").update({"transcribe_status": "error"}) \
                .eq("id", mid).execute()
            return jsonify({"status": "error", "error": r.get("error")})
        return jsonify({"status": status})
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Documents
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/documents", methods=["GET", "POST"])
@_require_auth
def api_documents():
    try:
        if request.method == "GET":
            q = supabase.table("mm_documents").select("*").eq("owner_id", OWNER_ID)
            if request.args.get("project_id"):
                q = q.eq("project_id", request.args["project_id"])
            if request.args.get("meeting_id"):
                q = q.eq("meeting_id", request.args["meeting_id"])
            r = q.order("uploaded_at", desc=True).execute()
            return jsonify(r.data or [])

        d = request.get_json() or {}
        file_b64 = d.get("file_b64", "")
        filename = d.get("filename", "document")
        if not file_b64:
            return jsonify({"ok": False, "error": "file_b64 required"}), 400
        raw = base64.b64decode(file_b64)
        path = f"docs/{OWNER_ID}/{datetime.datetime.utcnow().timestamp()}_{filename}"
        supabase.storage.from_(BUCKET).upload(
            path, raw, {"content-type": d.get("mime_type", "application/octet-stream")})
        url = supabase.storage.from_(BUCKET).get_public_url(path)
        row = {
            "owner_id":   OWNER_ID,
            "project_id": d.get("project_id"),
            "meeting_id": d.get("meeting_id"),
            "filename":   filename,
            "file_url":   url,
            "file_type":  d.get("mime_type"),
            "extracted_text": d.get("extracted_text"),
        }
        r = supabase.table("mm_documents").insert(row).execute()
        return jsonify(r.data[0] if r.data else {})
    except Exception as e:
        return err(e)


@mm_bp.route("/api/documents/<did>", methods=["DELETE"])
@_require_auth
def api_documents_delete(did):
    try:
        supabase.table("mm_documents").delete().eq("id", did).execute()
        return ok()
    except Exception as e:
        return err(e)


# ══════════════════════════════════════════════════════════════════
# Word export - matches the Pennyfarthing .dotx layout
# ══════════════════════════════════════════════════════════════════
@mm_bp.route("/api/export/<mid>.docx", methods=["GET"])
@_require_auth
def api_export_docx(mid):
    """Generate a .docx matching the house template: title block, attendees,
    optional Programme Review box, and the No./Item/Action/Date table."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from io import BytesIO

        m = supabase.table("mm_meetings").select("*").eq("id", mid).single().execute()
        meeting = m.data
        items = supabase.table("mm_meeting_items").select("*") \
            .eq("meeting_id", mid).order("sort_order").execute().data or []
        actions = supabase.table("mm_actions").select("*") \
            .eq("meeting_id", mid).execute().data or []
        tpl = {}
        if meeting.get("template_id"):
            t = supabase.table("mm_templates").select("*") \
                .eq("id", meeting["template_id"]).single().execute()
            tpl = t.data or {}
        proj = {}
        if meeting.get("project_id"):
            p = supabase.table("mm_projects").select("*") \
                .eq("id", meeting["project_id"]).single().execute()
            proj = p.data or {}
        people = {p["id"]: p for p in (supabase.table("mm_people")
                  .select("id,initials,name").execute().data or [])}

        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = sec.right_margin = Cm(1.8)
        sec.top_margin = sec.bottom_margin = Cm(1.5)

        # Header
        hdr = sec.header.paragraphs[0]
        hdr.text = tpl.get("doc_title") or "MEETING MINUTES"
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in hdr.runs:
            r.bold = True
            r.font.size = Pt(12)

        # Footer
        ftr = sec.footer.paragraphs[0]
        ftr.text = "Company Confidential"
        ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in ftr.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Title block
        date_str = meeting.get("meeting_date", "")
        try:
            date_str = datetime.date.fromisoformat(date_str).strftime("%d %B %Y")
        except Exception:
            pass
        p = doc.add_paragraph()
        run = p.add_run(date_str)
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        title_line = f'{proj.get("project_no","")} {proj.get("name","")}'.strip()
        if meeting.get("title"):
            title_line = f'{title_line} - {meeting["title"]}'
        run = p.add_run(title_line)
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph(f'Attendees: {meeting.get("attendees") or ""}')
        doc.add_paragraph(f'Apologies: {meeting.get("apologies") or ""}')

        # Programme Review box
        if tpl.get("has_programme_review"):
            doc.add_paragraph()
            pr = doc.add_table(rows=2, cols=4)
            pr.style = "Table Grid"
            labels = [
                ("Programme Review:", ""),
                ("Baseline Duration:", meeting.get("prog_baseline_duration") or ""),
                ("Prior Duration:", meeting.get("prog_prior_duration") or ""),
                ("Current Duration:", meeting.get("prog_current_duration") or ""),
            ]
            for i, (lab, val) in enumerate(labels):
                c = pr.rows[0].cells[i]
                c.text = ""
                para = c.paragraphs[0]
                rr = para.add_run(lab)
                rr.bold = True
                rr.font.size = Pt(8)
                if val:
                    para.add_run(f" {val}").font.size = Pt(8)
            row2 = [
                ("Baseline Completion:", meeting.get("prog_baseline_completion") or ""),
                ("", ""),
                ("Programme Anticipated Completion:",
                 meeting.get("prog_anticipated_completion") or ""),
                ("", ""),
            ]
            for i, (lab, val) in enumerate(row2):
                c = pr.rows[1].cells[i]
                c.text = ""
                para = c.paragraphs[0]
                if lab:
                    rr = para.add_run(lab)
                    rr.bold = True
                    rr.font.size = Pt(8)
                if val:
                    para.add_run(f" {val}").font.size = Pt(8)

        doc.add_paragraph()

        # Main minutes table
        has_actions = tpl.get("has_action_column", True)
        ncols = 4 if has_actions else 2
        table = doc.add_table(rows=1, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["No.", "Item", "Action", "Date"] if has_actions else ["No.", "Item"]
        for i, h in enumerate(headers):
            c = table.rows[0].cells[i]
            c.text = ""
            rr = c.paragraphs[0].add_run(h)
            rr.bold = True
            rr.font.size = Pt(9)

        acts_by_item = {}
        for a in actions:
            acts_by_item.setdefault(a.get("meeting_item_id"), []).append(a)

        for it in items:
            row = table.add_row()
            c0 = row.cells[0]
            c0.text = ""
            rr = c0.paragraphs[0].add_run(it["item_no"])
            rr.bold = True
            rr.font.size = Pt(9)

            c1 = row.cells[1]
            c1.text = ""
            para = c1.paragraphs[0]
            rr = para.add_run(it["title"])
            rr.bold = True
            rr.font.size = Pt(9)
            if it.get("content"):
                for line in it["content"].split("\n"):
                    if line.strip():
                        pp = c1.add_paragraph()
                        pr2 = pp.add_run(line.strip())
                        pr2.font.size = Pt(9)

            if has_actions:
                item_actions = acts_by_item.get(it["id"], [])
                who = ", ".join(
                    (people.get(a.get("assigned_to"), {}).get("initials")
                     or a.get("assigned_text") or "")
                    for a in item_actions if (a.get("assigned_to") or a.get("assigned_text"))
                )
                when = ", ".join(
                    a["due_date"] for a in item_actions if a.get("due_date"))
                c2 = row.cells[2]
                c2.text = ""
                c2.paragraphs[0].add_run(who).font.size = Pt(9)
                c3 = row.cells[3]
                c3.text = ""
                c3.paragraphs[0].add_run(when).font.size = Pt(9)

        doc.add_paragraph()
        cp = doc.add_paragraph()
        rr = cp.add_run("Circulation:")
        rr.bold = True
        rr.font.size = Pt(9)
        cp.add_run(" As Attendees and Apologies").font.size = Pt(9)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        fname = f'Minutes_{proj.get("project_no","")}_{meeting.get("meeting_date","")}.docx'
        return Response(
            buf.read(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{fname}"'},
        )
    except Exception as e:
        return err(e)
